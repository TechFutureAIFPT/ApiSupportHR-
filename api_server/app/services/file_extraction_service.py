from __future__ import annotations

import base64
from pathlib import Path
from io import BytesIO
from typing import Iterable, Literal

try:
    import fitz
except ModuleNotFoundError:  # pragma: no cover - optional in isolated test envs
    fitz = None  # type: ignore[assignment]

try:
    import pymupdf4llm
except ModuleNotFoundError:  # pragma: no cover - optional in isolated test envs
    pymupdf4llm = None  # type: ignore[assignment]

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover - optional in isolated test envs
    Document = None  # type: ignore[assignment]

from app.core.config import get_settings
from app.prompts import render_prompt
from app.services.gemini_service import generate_content

MIN_PDF_TEXT_LENGTH = 200
MAX_OCR_PAGES = 3
PDF_RENDER_SCALE = 2.2


def _normalize_text(text: str) -> str:
    return (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\t", " ")
        .replace(" ", " ")
        .strip()
    )


def _is_text_sufficient(text: str) -> bool:
    meaningful = "".join(ch for ch in text if ch.isalnum() or ch.isspace()).strip()
    return len(meaningful) >= MIN_PDF_TEXT_LENGTH


def _build_vision_prompt(document_type: Literal["cv", "jd"]) -> str:
    document_detail = (
        "Document type: Job Description. Prioritize title, requirements, skills, education, benefits, salary, and location."
        if document_type == "jd"
        else "Document type: Candidate CV. Prioritize name, email, phone, address, work history, education, skills, projects, and achievements."
    )
    return render_prompt(
        "file_extraction/ocr_document",
        context={"document_detail": document_detail},
    )


def _ocr_image_bytes(
    image_bytes: bytes,
    document_type: Literal["cv", "jd"],
    api_keys: Iterable[str] | None = None,
) -> str:
    settings = get_settings()
    response = generate_content(
        settings.gemini_default_model,
        [
            {
                "role": "user",
                "parts": [
                    {
                        "inlineData": {
                            "data": base64.b64encode(image_bytes).decode("utf-8"),
                            "mimeType": "image/png",
                        }
                    },
                    {"text": _build_vision_prompt(document_type)},
                ],
            }
        ],
        {
            "temperature": 0,
            "topP": 0.05,
            "topK": 1,
            "maxOutputTokens": 8192,
            "responseMimeType": "text/plain",
        },
        api_keys=api_keys,
    )
    return response.strip()


def _extract_text_from_pdf(
    file_bytes: bytes,
    force_ocr: bool,
    document_type: Literal["cv", "jd"],
    api_keys: Iterable[str] | None = None,
) -> str:
    if fitz is None or pymupdf4llm is None:
        raise RuntimeError("PDF extraction dependencies are not installed on server")

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        md_text = "" if force_ocr else pymupdf4llm.to_markdown(doc)
        if not force_ocr and _is_text_sufficient(md_text):
            return md_text

        pages = min(MAX_OCR_PAGES, doc.page_count)
        ocr_parts: list[str] = []
        for index in range(pages):
            page = doc.load_page(index)
            pix = page.get_pixmap(matrix=fitz.Matrix(PDF_RENDER_SCALE, PDF_RENDER_SCALE), alpha=False)
            ocr_text = _ocr_image_bytes(pix.tobytes("png"), document_type, api_keys=api_keys)
            if ocr_text:
                ocr_parts.append(ocr_text)

        combined_ocr = "\n\n".join(part for part in ocr_parts if part.strip())
        return combined_ocr or md_text
    finally:
        doc.close()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise RuntimeError("DOCX extraction dependency is not installed on server")

    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for line in cell.text.splitlines():
                    line = line.strip()
                    if line and (not table_lines or table_lines[-1] != line):
                        table_lines.append(line)

    return "\n".join(paragraphs + table_lines)


def _decode_plain_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_upload(
    file_bytes: bytes,
    filename: str,
    content_type: str,
    force_ocr: bool = False,
    document_type: str | None = None,
    api_keys: Iterable[str] | None = None,
) -> str:
    settings = get_settings()
    max_size_mb = settings.upload_file_size_limit_mb
    if len(file_bytes) > max_size_mb * 1024 * 1024:
        raise ValueError(f"File is too large. Maximum size is {max_size_mb}MB.")

    name_lower = filename.lower()
    suffix = Path(name_lower).suffix.lower()
    allowed_extensions = set(settings.allowed_upload_extensions)
    allowed_mime_types = set(settings.allowed_upload_mime_types)
    if suffix and suffix not in allowed_extensions:
        raise ValueError(f"Unsupported file extension: {suffix}")
    normalized_content_type = (content_type or "").strip().lower()
    if normalized_content_type and normalized_content_type not in allowed_mime_types and not normalized_content_type.startswith("image/"):
        raise ValueError(f"Unsupported MIME type: {content_type}")

    doc_type: Literal["cv", "jd"] = "jd" if document_type == "jd" else "cv"

    if content_type == "application/pdf" or name_lower.endswith(".pdf"):
        raw_text = _extract_text_from_pdf(file_bytes, force_ocr, doc_type, api_keys=api_keys)
    elif (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or name_lower.endswith(".docx")
    ):
        raw_text = _extract_text_from_docx(file_bytes)
    elif content_type.startswith("image/"):
        raw_text = _ocr_image_bytes(file_bytes, doc_type, api_keys=api_keys)
    elif content_type in {"text/plain", "text/csv", "application/csv"} or name_lower.endswith((".txt", ".csv")):
        raw_text = _decode_plain_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}")

    normalized = _normalize_text(raw_text)
    if not normalized:
        raise ValueError(f"Could not extract text from file: {filename}")
    return normalized
